import base64
from io import BytesIO
from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell


def _replace_image_run(paragraph, run, image_bytes: bytes):
    """
    Replaces the given run in a paragraph with an image.
    """
    paragraph._element.remove(run._element)
    new_run = paragraph.add_run()
    new_run.add_picture(BytesIO(image_bytes))


def replace_placeholders(doc_or_cell: _Document | _Cell, data: dict) -> None:
    """
    Replaces all text placeholders and images (by alt text) in a Word document or table cell.

    Args:
        doc_or_cell (_Document | _Cell): The docx Document or table cell to process.
        data (dict): A dictionary with keys matching placeholders or image alt texts,
                     and values being either text or base64-encoded images.
    """
    for p in doc_or_cell.paragraphs:
        # Replace text placeholders
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

        # Replace images by alt text
        for run in p.runs:
            if run._element.xpath('.//*[local-name()="pic"]'):
                drawing = run._element.xpath('.//*[local-name()="docPr"]')
                if drawing:
                    alt_text = drawing[0].get('descr')
                    if alt_text:
                        key = alt_text.strip("[]#")
                        if key in data:
                            image_data = base64.b64decode(data[key])
                            _replace_image_run(p, run, image_data)
                            break

    for table in doc_or_cell.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell, data)


def render_docx(template_path: str, output_path: str, data: dict) -> None:
    """
    Loads a DOCX template, replaces placeholders, and saves the result.

    Args:
        template_path (str): Path to the input DOCX template.
        output_path (str): Path to save the generated DOCX file.
        data (dict): Dictionary with text and base64 image data.
    """
    doc = Document(template_path)
    replace_placeholders(doc, data)
    doc.save(output_path)
